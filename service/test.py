from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table, _Row, _Cell

from docx.styles.style import _ParagraphStyle,_CharacterStyle

import re
from copy import deepcopy

S_REPLACE:list[str] = ["{{","}}"]
S_FORMULA:list[str] = ["{%","%}"]
S_TABLE:str = S_FORMULA[0]+'TABLE'+S_FORMULA[1]

DEMO:dict = {
  "organisation": "Orga",
  "vertreten_durch": "Vertreter",
  "strasse": "Musterstr",
  "hausnummer":"1",
  "plz": "30167",
  "ort": "Hannover",
  "d_angebot": "01.07.2022",
  "d_guetlig_bis": "02.07.2022",
  "d_nutzung_start": "03.07.2022",
  "d_nutzung_end": "04.07.2022",
  "p_reinigung": 1,
  "p_headphone": 25,
  "p_kabel": 5,
  "p_sender": 200,
  "order": [
    {
      "pos":1,
      "text":"Kopfhörer",
      "subtext": "",
      "menge": 100,
      "p_einzel": 1.8,
      "summe": 180
    },
    {
      "pos":2,
      "text":"Sender",
      "subtext": "",
      "menge": 3,
      "p_einzel": 30,
      "summe": 90
    },
    {
      "pos":3,
      "text":"Aufwandspauschale",
      "subtext": "(Laden und Reinigen der KH nach der Vermietung, Ausgabe und Annahme der vermieteten Gegenstände)",
      "menge": 1,
      "p_einzel": 50,
      "summe": 50
    }
  ]
}

def btwn(text,sep:list[str]=S_REPLACE):
  return f'{sep[0]}{text}{sep[1]}'

def replace_text(text:str, backbone:dict=DEMO, parentKey:str=None, enumeration:int=None) -> str:
  if S_REPLACE[0] not in text and S_REPLACE[1] not in text:
    return text
  for key, value in backbone.items():
    if key in text:
      text_types = [str,int]
      if type(value) in text_types:
        text = text.replace(btwn(key), str(value))
        if parentKey:
          if enumeration:
            text = text.replace(btwn(f'{parentKey}[{enumeration}].{key}'), str(value))
          else:
            text = text.replace(btwn(f'{parentKey}.{key}'), str(value))
      elif type(value) == list:
        for i,e in enumerate(value):
          text = replace_text(text, e, parentKey=key, enumeration=i+1)
      elif type(value) == dict:
        text = replace_text(text, e, parentKey=key)
  return text

def replace_in_paragraph(p:Paragraph) -> Paragraph:
  runs:list[Run] = p.runs
  for run in runs:
    text = replace_text(run.text)
    run.text = text
  return p


def iterate_table(table:Table):
  def remove_row(table:Table, row:_Row | None):
    if row:
      tbl = table._tbl
      tr = row._tr
      tbl.remove(tr)
  
  key:str=None
  iteration:list=None
  
  first_cell:_Cell = table.rows[0].cells[0]
  fc_text:str = first_cell.text
  
  params = []
  
  for k, v in DEMO.items():
    if f'{S_REPLACE[0]}{k}' in fc_text:
      key, iteration = k, v
      break
  if not key:
    return table
  
  if '|' in fc_text:
    start: int = fc_text.find('|')+1
    end: int = fc_text.find(S_REPLACE[1])
    param_s = fc_text[start:end].replace(' ','')
    params = param_s.split(',')
  
  headline:_Row|None = table.rows[1] if 'headline' in params else None
  workRow_number:int = 2 if headline else 1
  workRow:_Row = table.rows[workRow_number]
  
  for num,e in enumerate(iteration):
    newRow:_Row = table.add_row()
    for w_cell,n_cell in zip(workRow.cells,newRow.cells):
      for w_p in w_cell.paragraphs:
        n_p:Paragraph = n_cell.add_paragraph()
        n_p.style= w_p.style
        w_runs:list[Run] = w_p.runs
        for w_run in w_runs:
          text: str = w_run.text
          for k,v in e.items():
            text = text.replace(S_REPLACE[0]+k+S_REPLACE[1],f'{S_REPLACE[0]}{key}[{num+1}].{k}{S_REPLACE[1]}')
          n_p.add_run(text)
          
  remove_row(table, table.rows[workRow_number])
  remove_row(table, table.rows[0])
  return table
    

doc = Document('./templates/angebot.docx')



for t in doc.tables:
  t = iterate_table(t)


for t in doc.tables:
  for row in t.rows:
    for cell in row.cells:
      for p in cell.paragraphs:
        p = replace_in_paragraph(p)

for p in doc.paragraphs:
  p = replace_in_paragraph(p)
  

doc.save('./temporary/test.docx')