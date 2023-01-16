from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table, _Row, _Cell
from docx.styles.style import _ParagraphStyle,_CharacterStyle
from copy import deepcopy


class Docifyer():
  
  S_REPLACE:list[str] = ["{{","}}"]
  S_FORMULA:list[str] = ["{%","%}"]
  S_TABLE:str = S_FORMULA[0]+'TABLE'+S_FORMULA[1]
  
  def __init__(self, name:str, data:dict, template_path:str, temporary_path:str):
    self.name = name
    self.template_path = template_path
    self.temporary_path = temporary_path
    self.doc = Document(self.template_path+f'/{self.name}.docx')
    self.data = data
    
  def run(self):
    for t in self.doc.tables:
      t = self._iterate_table(t)
    
    for t in self.doc.tables:
      for row in t.rows:
        for cell in row.cells:
          for p in cell.paragraphs:
            p = self._replace_in_paragraph(p)
    
    for p in self.doc.paragraphs:
      p = self._replace_in_paragraph(p)
  
  def save(self, path:str = None, thema:str= None, date:str = None) -> str:
    path = path or self.temporary_path
    name: str = self.name
    if thema:
      name += f'_{thema}'
    if date:
      name += f'_{date}'
    name += '.docx'
    full_path = path+f'/{name}'
    self.doc.save(full_path)
    return name

  def _btwn(self, text,sep:list[str]=None):
    sep = sep or self.S_REPLACE
    return f'{sep[0]}{text}{sep[1]}'
  
  def _to_str(self, text):
    if type(text) == float:
      return f'{text:.2f}'
    return str(text)

  def _replace_text(self, text:str, backbone:dict, parentKey:str=None, enumeration:int=None) -> str:
    if self.S_REPLACE[0] not in text and self.S_REPLACE[1] not in text:
      return text
    value: str | int | float | list | dict
    for key, value in backbone.items():
      if key in text: #TODO: pipe-operator einbauen!
        text_types = [str,int,float]
        if type(value) in text_types:
          text = text.replace(self._btwn(key), self._to_str(value))
          if parentKey:
            if enumeration:
              text = text.replace(self._btwn(f'{parentKey}[{enumeration}].{key}'), self._to_str(value))
            else:
              text = text.replace(self._btwn(f'{parentKey}.{key}'), self._to_str(value))
        elif type(value) == list:
          for i,e in enumerate(value):
            text = self._replace_text(text, e, parentKey=key, enumeration=i+1)
        elif type(value) == dict:
          text = self._replace_text(text, e, parentKey=key)
    return text

  def _replace_in_paragraph(self, p:Paragraph) -> Paragraph:
    runs:list[Run] = p.runs
    for run in runs:
      text = self._replace_text(run.text, self.data)
      run.text = text
    return p

  def _get_options(self, text:str) -> list | None: #TODO: einbauen!
    if '|' not in text:
      return None
    start: int = text.find('|')+1
    end: int = text.find(self.S_REPLACE[1])
    param_s = text[start:end].replace(' ','')
    params = param_s.split(',')
    return params

  def _iterate_table(self, table:Table):
    def remove_row(table:Table, row:_Row | None):
      if row:
        tbl = table._tbl
        tr = row._tr
        tbl.remove(tr)
    
    key:str=None
    iteration:list=None
    
    first_cell:_Cell = table.rows[0].cells[0]
    fc_text:str = first_cell.text
    
    params = []
    
    for k, v in self.data.items():
      if f'{self.S_REPLACE[0]}{k}' in fc_text:
        key, iteration = k, v
        break
    if not key:
      return table
    
    if '|' in fc_text:
      start: int = fc_text.find('|')+1
      end: int = fc_text.find(self.S_REPLACE[1])
      param_s = fc_text[start:end].replace(' ','')
      params = param_s.split(',')
    
    workRow_number:int = 1
    footline_number:int = 2
    if 'headline' in params:
      workRow_number += 1
      footline_number += 1
    workRow:_Row = table.rows[workRow_number]
    
    
    for num,e in enumerate(iteration):
      newRow:_Row = table.add_row()
      for w_cell,n_cell in zip(workRow.cells,newRow.cells):
        for w_p in w_cell.paragraphs:
          n_p:Paragraph = n_cell.add_paragraph()
          n_p.style= w_p.style
          w_runs:list[Run] = w_p.runs
          for w_run in w_runs:
            text: str = w_run.text
            for k,v in e.items():
              text = text.replace(self.S_REPLACE[0]+k+self.S_REPLACE[1],f'{self.S_REPLACE[0]}{key}[{num+1}].{k}{self.S_REPLACE[1]}')
            n_p.add_run(text)
    
    remove_row(table, table.rows[workRow_number])
    remove_row(table, table.rows[0])
    if "footline" in params:
      footline:_Row = table.rows[footline_number]
      new_footline = table.add_row()
      new_footline = footline
    return table
